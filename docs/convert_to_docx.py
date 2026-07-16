"""Convert permission markdown files to formatted .docx files."""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Brand colours
GREEN_DARK  = RGBColor(0x1F, 0x57, 0x3B)   # dark green headings
GREEN_MID   = RGBColor(0x2D, 0x7A, 0x55)   # table header bg  (set via XML)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
GREY_LIGHT  = RGBColor(0xF4, 0xF4, 0xF4)   # alt table row
BLACK       = RGBColor(0x1A, 0x1A, 0x1A)


def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_border(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{'insideH' if side in ('top','bottom') else side}")
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:color"), "CCCCCC")
        borders.append(el)
    tcPr.append(borders)


def apply_run_inline(run, text: str):
    """Apply bold/italic/inline-code formatting within a run."""
    run.text = text


def add_inline_text(para, text: str):
    """Add text to paragraph, honouring **bold** and `code` markers."""
    # Split on bold (**text**) and inline code (`text`)
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
            run.font.color.rgb = BLACK
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name   = "Courier New"
            run.font.size   = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            if part:
                run = para.add_run(part)
                run.font.color.rgb = BLACK


def style_heading(para, level: int):
    para.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    para.paragraph_format.space_after  = Pt(4)
    run = para.runs[0] if para.runs else para.add_run()
    run.font.color.rgb = GREEN_DARK
    run.font.bold      = True
    run.font.size      = Pt({1: 20, 2: 15, 3: 12}.get(level, 11))


def convert(md_path: str, docx_path: str):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.8)

    # Default body style
    style = doc.styles["Normal"]
    style.font.name  = "Calibri"
    style.font.size  = Pt(10.5)
    style.font.color.rgb = BLACK

    lines      = open(md_path, encoding="utf-8").readlines()
    i          = 0
    in_code    = False
    code_lines = []

    while i < len(lines):
        raw  = lines[i].rstrip("\n")
        line = raw.strip()

        # ── fenced code block ────────────────────────────────────────────────
        if line.startswith("```"):
            if not in_code:
                in_code    = True
                code_lines = []
            else:
                in_code = False
                para    = doc.add_paragraph()
                para.paragraph_format.space_before = Pt(4)
                para.paragraph_format.space_after  = Pt(4)
                para.paragraph_format.left_indent  = Cm(0.8)
                run = para.add_run("\n".join(code_lines))
                run.font.name  = "Courier New"
                run.font.size  = Pt(8.5)
                run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)
            i += 1
            continue

        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        # ── horizontal rule ──────────────────────────────────────────────────
        if line in ("---", "***", "___") or re.match(r"^-{3,}$", line):
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)
            run  = para.add_run()
            pPr  = para._p.get_or_add_pPr()
            pb   = OxmlElement("w:pBdr")
            bot  = OxmlElement("w:bottom")
            bot.set(qn("w:val"),   "single")
            bot.set(qn("w:sz"),    "6")
            bot.set(qn("w:space"), "1")
            bot.set(qn("w:color"), "2D7A55")
            pb.append(bot)
            pPr.append(pb)
            i += 1
            continue

        # ── headings ─────────────────────────────────────────────────────────
        m = re.match(r"^(#{1,3})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text  = m.group(2)
            hmap  = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
            para  = doc.add_paragraph(style=hmap.get(level, "Heading 3"))
            para.clear()
            run = para.add_run(text)
            style_heading(para, level)
            i += 1
            continue

        # ── blockquote / note ────────────────────────────────────────────────
        if line.startswith("> "):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent  = Cm(1)
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after  = Pt(4)
            run = para.add_run(line[2:])
            run.italic         = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # ── table ────────────────────────────────────────────────────────────
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1

            # Remove separator row (---|---)
            rows = [r for r in table_lines
                    if not re.match(r"^\|[-| :]+\|$", r)]
            if not rows:
                continue

            # Parse cells
            parsed = []
            for r in rows:
                cells = [c.strip() for c in r.strip("|").split("|")]
                parsed.append(cells)

            cols    = max(len(r) for r in parsed)
            tbl     = doc.add_table(rows=len(parsed), cols=cols)
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            tbl.style     = "Table Grid"

            for ri, row_data in enumerate(parsed):
                for ci, cell_text in enumerate(row_data):
                    if ci >= cols:
                        break
                    cell = tbl.cell(ri, ci)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    set_cell_border(cell)

                    if ri == 0:
                        set_cell_bg(cell, "1F573B")
                        p   = cell.paragraphs[0]
                        run = p.add_run(re.sub(r"\*\*(.+?)\*\*", r"\1", cell_text))
                        run.bold           = True
                        run.font.color.rgb = WHITE
                        run.font.size      = Pt(9.5)
                        p.paragraph_format.space_before = Pt(3)
                        p.paragraph_format.space_after  = Pt(3)
                    else:
                        bg = "F4F9F6" if ri % 2 == 0 else "FFFFFF"
                        set_cell_bg(cell, bg)
                        p = cell.paragraphs[0]
                        add_inline_text(p, cell_text)
                        for run in p.runs:
                            run.font.size = Pt(9.5)
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.space_after  = Pt(2)

            doc.add_paragraph()  # spacing after table
            continue

        # ── checklist item ───────────────────────────────────────────────────
        if re.match(r"^- \[[ x]\]", line):
            checked = line[3] == "x"
            text    = line[6:].strip()
            para    = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.left_indent  = Cm(0.5)
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after  = Pt(1)
            prefix  = "[x] " if checked else "[ ] "
            run     = para.add_run(prefix)
            run.font.name      = "Courier New"
            run.font.size      = Pt(10)
            run.font.color.rgb = GREEN_DARK if checked else RGBColor(0x88, 0x88, 0x88)
            add_inline_text(para, text)
            i += 1
            continue

        # ── bullet list ──────────────────────────────────────────────────────
        if re.match(r"^[-*]\s+", line):
            text = re.sub(r"^[-*]\s+", "", line)
            para = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.left_indent  = Cm(0.5)
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after  = Pt(1)
            add_inline_text(para, text)
            i += 1
            continue

        # ── blank line ───────────────────────────────────────────────────────
        if not line:
            i += 1
            continue

        # ── normal paragraph ─────────────────────────────────────────────────
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after  = Pt(4)
        add_inline_text(para, line)
        i += 1

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == "__main__":
    base = r"c:\Users\vinukiT\Downloads\JAS\docs"

    convert(
        f"{base}\\permission_business_case.md",
        f"{base}\\permission_business_case.docx",
    )
    convert(
        f"{base}\\permission_technical.md",
        f"{base}\\permission_technical.docx",
    )
